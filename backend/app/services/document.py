"""文档服务：负责加载不同格式的文档，并按「标题/章节结构」智能切分"""
import re
from pathlib import Path

from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter


# 识别「标题行」的正则（匹配行首）
_HEADER_PATTERNS = [
    re.compile(r"^#{1,6}\s+\S+"),         # Markdown 标题：## 账号与密码
    re.compile(r"^[一二三四五六七八九十百]+、"),  # 中文编号：一、二、三、
    re.compile(r"^（[一二三四五六七八九十百]+）"),  # 括号编号：（一）（二）
    re.compile(r"^\d+[.、]\s*\S+"),         # 数字编号：1. / 1、
]


def _is_header(line: str) -> bool:
    """判断一行是不是「标题行」"""
    line = line.strip()
    if not line:
        return False
    return any(p.match(line) for p in _HEADER_PATTERNS)


def load_document(file_path: str | Path) -> str:
    """根据文件后缀，读取文档的纯文本内容"""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix in (".txt", ".md"):
        # 纯文本文件：直接读取
        return path.read_text(encoding="utf-8")

    if suffix == ".pdf":
        # PDF：用 pypdf 逐页提取文字
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    if suffix == ".docx":
        # Word：用 python-docx 逐段提取文字
        import docx

        doc = docx.Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)

    raise ValueError(f"不支持的文件类型：{suffix}")


def split_by_structure(text: str) -> list[tuple[str, str]]:
    """按标题把文档切成「章节」，返回 [(章节标题, 正文), ...]

    比纯字符切分更聪明：把「标题行」当作切分边界，
    一个章节的正文保持完整，不会被生硬截断。
    """
    sections: list[tuple[str, str]] = []
    current_header = ""            # 当前章节标题（空 = 还没遇到标题的开头部分）
    current_body: list[str] = []

    for line in text.split("\n"):
        if _is_header(line):
            # 遇到新标题：先把上一个章节收尾，再开始新的
            if current_body:
                sections.append((current_header, "\n".join(current_body)))
            current_header = line.strip()
            current_body = []
        else:
            current_body.append(line)

    # 收尾：最后一个章节
    if current_body:
        sections.append((current_header, "\n".join(current_body)))
    return sections


def process_document(
    file_path: str | Path,
    chunk_size: int = 300,
    chunk_overlap: int = 20,
) -> list[Document]:
    """加载文档并按「标题结构」智能切分（RAG 的「索引」环节）

    相比纯字符切分，做了两件事：
    1. 按标题/章节切分：把标题行当作边界，保证一个章节不被生硬截断
    2. 标题继承：把章节标题注入到每个块里，避免「上下文丢失」
       （检索到的内容自带「它属于哪个章节」，语义更完整）

    参数：
        chunk_size：单个文本块的最大长度（只有超长章节才会被进一步按句子切分）
        chunk_overlap：相邻块之间的重叠长度
    返回：
        list[Document]：每个 Document 是一段文本 + 元信息（来源文件名、章节标题）
    """
    text = load_document(file_path)
    source_name = Path(file_path).name

    # 章节内再切分：当某个章节正文太长时，按句子切，保证块不会过长
    sentence_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["。", "！", "？", "；", "，", " ", ""],
    )

    docs: list[Document] = []
    for header, body in split_by_structure(text):
        body = body.strip()
        if not body:
            continue

        if header:
            # 标题继承：正文较短则整节一块；过长则按句子切，但每块都带标题前缀
            if len(header) + len(body) + 1 <= chunk_size:
                pieces = [f"{header}\n{body}"]
            else:
                pieces = [f"{header}\n{p}" for p in sentence_splitter.split_text(body)]
        else:
            # 没有标题的开头部分（如文档标题），按句子切
            pieces = sentence_splitter.split_text(body)

        for piece in pieces:
            docs.append(
                Document(
                    page_content=piece,
                    metadata={"source": source_name, "section": header},
                )
            )

    return docs
